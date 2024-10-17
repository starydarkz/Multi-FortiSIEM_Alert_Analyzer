#Multi-Fortisiem Alert Analyzer v3.0  | Created By: StaryDarkz  | Telegram: https://t.me/StaryDarkz 

'''
Novedades de la nueva version

~Top Graficos de alertas por instancia de SIEM
~Analisis de alertas general por hora
~Analisis de recurrencia de alertas de incidentes
~Mejora en el archivo de configuracion
~Nuevo logo del proyecto
~Mejor documentacion y organizacion de codigo
'''

import docx, time, datetime, re, httplib2
import matplotlib.pyplot as plt
from docx.enum.text import WD_ALIGN_PARAGRAPH 
from xml.dom.minidom import Node, parseString
from Resources.config import allsiem
import xml.etree.ElementTree as ET
from os import sys


#Funciones
def add_dic_data(dic, key, value):
    if key in dic:
        dic[key].append(value)
    else:
        dic[key] = [value]
    return dic

#Funciones DOCX
def add_table(plantilla, data): 

    table = plantilla.add_table(rows = 0, cols=2, style = "Grid Table 4 - Accent 11")          
    row = table.add_row().cells 

    row[0].text = "Rule Name"
    row[1].text = "Total Alertas"

    for incidente in data.keys():
        row = table.add_row().cells
        row[0].text = incidente
        row[1].text = str(data[incidente])

def graficar_gbarras(data_json, gname):
    """ Grafico de barras"""
    keys = []
    values = []

    for element in data_json:
        keys.append(element)
        values.append(data_json[element])

    #Graficar solo el TOP 24
    if len(keys) > 24:
        keys = keys[0:10]
        values = values[0:10]


    plt.rcParams.update(plt.rcParamsDefault) #Reinicia Estilo
    #plt.style.use('seaborn') #Aplica Estilo
    fig, ax = plt.subplots(figsize=(7, 4)) #Aplica Tamaño Figura
    ax.barh(keys, values, height=0.8) #Agrega Datos Ejes
    ax.invert_yaxis() #Invierte Impresión de Valores
    ax.xaxis.grid(False); ax.yaxis.grid(False) #Elimina Ejes Fondo Figura

        #Genera Etiquetas (Valores) Lado Derecha Barras
    for i in range(len(values)):
        ax.text(values[i]+1, i, values[::][i], color='black', fontsize=8, ha='center',  fontweight='bold', bbox=dict(facecolor='skyblue', alpha=0.3))

    plt.tick_params(axis='x', which='both', bottom=False, top=False, labelbottom=False) #Elimina Datos Eje X
    plt.yticks(fontsize=8) #Etiquetas Eje Y
    fig.tight_layout()

    #Save Image
    name = f"{gname}.png"
    plt.savefig(name)
    plt.close()

def graficar_gbarras_CLIENTES(data_json, gname):
    """ Grafico de barras"""
    keys = []
    values = []

    for element in data_json:
        keys.append(element)
        values.append(data_json[element])

    #Graficar solo el TOP 10
    if len(keys) > 10:
        keys = keys[0:10]
        values = values[0:10]


    plt.rcParams.update(plt.rcParamsDefault) #Reinicia Estilo
    #plt.style.use('seaborn') #Aplica Estilo
    fig, ax = plt.subplots(figsize=(7, 4)) #Aplica Tamaño Figura
    ax.barh(keys, values, height=0.8) #Agrega Datos Ejes
    ax.invert_yaxis() #Invierte Impresión de Valores
    ax.xaxis.grid(False); ax.yaxis.grid(False) #Elimina Ejes Fondo Figura

        #Genera Etiquetas (Valores) Lado Derecha Barras
    for i in range(len(values)):
        ax.text(values[i]+1, i, values[::][i], color='black', fontsize=8, ha='center',  fontweight='bold', bbox=dict(facecolor='skyblue', alpha=0.3))

    plt.tick_params(axis='x', which='both', bottom=False, top=False, labelbottom=False) #Elimina Datos Eje X
    plt.yticks(fontsize=8) #Etiquetas Eje Y
    fig.tight_layout()

    #Save Image
    name = f"{gname}.png"
    plt.savefig(name)
    plt.close()


#Funciones de la API FortiSIEM
def select_query(xmlquery, input_time_user, email_alert):
    
    timestamp = int(time.time())

    custom_time = 86400 * input_time_user

    time_start = timestamp - custom_time
    time_end = timestamp
    xml_incident_count = f"""<?xml version="1.0" encoding="UTF-8"?>
    <Reports>
    <Report baseline="" rsSync="">
        <Name>Incident Notification Count</Name>
        <Description> Total de Notificaciones con email alert especificado</Description>
        <CustomerScope groupByEachCustomer="false">
        </CustomerScope>
        <SelectClause>
            <AttrList>ruleName,incidentId,phRecvTime,incidentSrc,incidentTarget,incidentDetail</AttrList>
        </SelectClause>
        <OrderByClause>
            <AttrList>ruleName DESC</AttrList>
        </OrderByClause>
        <PatternClause window="3600">
            <SubPattern id="1164394" name="Filter_OVERALL_STATUS">
                <SingleEvtConstr>(phEventCategory = 3  AND  eventType = "PH_INCIDENT_ACTION_STATUS" AND actionName CONTAIN "{email_alert}")</SingleEvtConstr>
            </SubPattern>
        </PatternClause>
        <userRoles>
            <roles custId="2001">1169250</roles>
        </userRoles>
        <SyncOrgs/><ReportInterval>
            <Low>{time_start}</Low>
            <High>{time_end}</High>error
        </ReportInterval>
        <TrendInterval>auto</TrendInterval>
        <TimeZone/>
    </Report>
    </Reports>"""
 
    if "xml_incident_count" == xmlquery:
        return xml_incident_count

def get_queryfromsiem(ip_siem, user, password, xml_query):

    url = "https://" + ip_siem + ":443/phoenix/rest/query/"
    urlfirst = url + "eventQuery"
    
    h = httplib2.Http(disable_ssl_certificate_validation=True)
    h.add_credentials(user, password)
    
    header = {'Content-Type': 'text/xml'}
    
    doc = parseString(xml_query)
    t = doc.toxml()

    if '<DataRequest' in t:
        t1 = t.replace("<DataRequest", "<Reports><Report")
    else:
        t1 = t
    if '</DataRequest>' in t1:
        t2 = t1.replace("</DataRequest>", "</Report></Reports>")
    else:
        t2 = t1

    resp, content = h.request(urlfirst, "POST", t2, header)
    
    #print (resp, content)
    
    queryId = content.decode("utf-8")
    if "xml version" in queryId:
        queryId = extrat_data_query(queryId)
        
    if 'error code="255"' in queryId:
        print ("Query Error, check sending XML file.")
        #exit()
        return "Error"

    urlSecond = url + "progress/" + queryId
    if resp['status'] == '200':
        resp, content = h.request(urlSecond)

    else:
        print ("appServer doesn't return query. Error code is %s" % resp['status'] )
        #exit()
        return "Error"


    while True:
        resp, content = h.request(urlSecond)
        try: 
            progreso = extrat_data_status(content)
            #print (progreso)
            if progreso == 100:
                break
        except:
            while content.decode("utf-8") != '100':
                resp, content = h.request(urlSecond)
            break

    outXML = []
   
    
    urlFinal = url + 'events/' + queryId + '/0/1000'
    resp, content = h.request(urlFinal)
    
    #print (resp, "\n\n", content)

    if content != '':
        outXML.append(content.decode("utf-8"))

    p = re.compile('totalCount="\d+"')
    mlist = p.findall(content.decode())
    
    
    if mlist[0] != '':
        mm = mlist[0].replace('"', '')
        m = mm.split("=")[-1]
        num = 0
        if int(m) > 1000:
            result = []
            
            num = int(m) / 1000
            num = int(num)
            if int(m) % 1000 > 0:
                num += 1
        if num > 0:

            for i in range(num):
                urlFinal = url + 'events/' + queryId + '/' + str((i * 1000)+1) + '/1000'
                resp, content = h.request(urlFinal)
                if content != '':
                    outXML.append(content.decode("utf-8"))   
                    data = dumpXML(outXML)
                    result.extend(data)
            return result             
    else:
        print ("no info in this report.")
        return "Error"
    data = dumpXML(outXML)
    return data

def extrat_data_status(xml_string):
    query = ET.fromstring(xml_string).find('./result/progress').text
    return int(query)

def extrat_data_query(xml):
    query = (ET.fromstring(xml).get('requestId'))
    timestamp = ET.fromstring(xml).find('./result/expireTime').text
    resultado = f"{query},{timestamp}"
    return resultado

def dumpXML(xmlList):
    param = []
    for xml in xmlList:
        doc = parseString(xml.encode('ascii', 'xmlcharrefreplace'))
    for node in doc.getElementsByTagName("events"):
        for node1 in node.getElementsByTagName("event"):
            mapping = {}
            for node2 in node1.getElementsByTagName("attributes"):
                for node3 in node2.getElementsByTagName("attribute"):
                    itemName = node3.getAttribute("name")
                    for node4 in node3.childNodes:
                        if node4.nodeType == Node.TEXT_NODE:
                            message = node4.data
                            if '\n' in message:
                                message = message.replace('\n', '')
                            mapping[itemName] = message
            param.append(mapping)
    return param


#Funciones de analisis
def generate_eventcount(param, element):
    if len(param) == 0:
        print (f"No records found on {element}")
        return "Error"
    else:
        dicdata = {}

        #Limpiar updates [FALTA]
        incidentsId = []
        incidents= []
        count = 0
        for element in param:
            count += 1
            if element["incidentId"] not in incidentsId:
                incidentsId.append(element["incidentId"]) 
                incidents.append(element)           
        

        for item in incidents: #Generar Count de incidentes unicos
            if item["ruleName"] not in dicdata.keys():
                dicdata[item["ruleName"]] = 1
            else:
                dicdata[item["ruleName"]] += 1
        return dicdata, len(incidentsId), incidents

def extrat_keyandvalue_CLIENTS(data):

    datanew = {}

    for cliente in data:
        totalalerts = 0
        for rulenamevalue in data[cliente]:
            totalalerts += data[cliente][rulenamevalue]
        
        datanew[cliente] = totalalerts
    return datanew

def get_time_for_reportname():

    fecha = datetime.datetime.now()
    fecha = fecha.strftime("%Y_%m_%d %H_%M") 

    return fecha

def extrat_keyandvalue(data):
    dataend = {}

    for element in data:
        dataend[element] = data[element]
    return dataend
   
def rule_pattern_analyzer(incidents):

    incidents_with_details = []
    incidents_with_details_data = {}
    end_dic_data = {}

    for incident in incidents: # Extraer incidentes con details
        
        try:
            a = (incident["incidentDetail"])
            if (incident["ruleName"]) not in incidents_with_details:
                incidents_with_details.append(incident["ruleName"])
        except:
            pass

    for incident in incidents: # Agrupar details por incidentes
        #print (element["ruleName"], incidents_with_details)
        if (incident["ruleName"]) in incidents_with_details:
            incidents_with_details_data = add_dic_data(incidents_with_details_data, (incident["ruleName"]), (incident["incidentDetail"]))
    
    for rule in incidents_with_details_data:
        #rule
        detail_count = {}
        for element in incidents_with_details_data[rule]:
            if element not in detail_count.keys():
                detail_count[element] = 1
            else:
                detail_count[element] += 1


        end_dic_data[rule] = detail_count
    
    # for element in end_dic_data.keys():
    #     print (element)
    #     print (input())
    #     print (end_dic_data[element])
        
    #     print (input())

    
    max_values = {}

    for atributo, subatributos in end_dic_data.items():
        # Encuentra el subatributo con el valor máximo
        subatributo_max = max(subatributos, key=subatributos.get)
        if subatributos[subatributo_max] > 10:
            max_values[atributo] = {subatributo_max: subatributos[subatributo_max]}

    return (max_values)

def count_by_hour(all_incidents):

    data_by_hours = {
        "00":0,
        "01":0,
        "02":0,
        "03":0,
        "04":0,
        "05":0,
        "06":0,
        "07":0,
        "08":0,
        "09":0,
        "10":0,
        "11":0,
        "12":0,
        "13":0,
        "14":0,
        "15":0,
        "16":0,
        "16":0,
        "17":0,
        "18":0,
        "19":0,
        "20":0,
        "21":0,
        "22":0,
        "23":0
    }
    for siem in all_incidents:

        for element in all_incidents[siem]:
            hour = re.findall('\D*\s\D*\d*\d\s(?P<Hour>\d*):\d*:\d*', element["phRecvTime"])
            data_by_hours[hour[0]] += 1
    
    return data_by_hours
            

#Funciones Generales
def menu(select):
    import getpass

    if select == "login":
        print ("\nIngresa tus credenciales:\n")
        username = input("Username:\n-->")
        password = getpass.getpass("Password:\n-->")
        return username, password

    elif select == "select_time":
        print ("Especifica el rango de tiempo del reporte:\n\n")
        time = int(input("Cantidad de Dias hacia atras\n-->"))
        return time

    elif select == "email_alert":
        print ("Especifica el correo de notificaciones:")
        email = input("Email: -->") 

def main(allsiem = allsiem):
    """ Ejecucion inicial """

    #Datos de entrada requeridos
    username, password = menu("login")
    time = menu("select_time")
    email_alert = menu("email_alert")

    #Cargar la plantilla
    plantilla = docx.Document("Resources/plantillaMFAA.docx")
    plantilla.add_page_break()

    #Iteracion de cada cliente
    alldata = {}
    all_count_bysiem = {}
    totalalerts = 0
    all_count_alerts_by_siem = {}
    all_alerts_general_count = {}

    all_patterns= {}

    for element in allsiem:
        """ Iteracion por cada cliente"""

        data = get_queryfromsiem(allsiem[element], f"super/{username}", password, select_query("xml_incident_count", time, email_alert))

        try:
            all_count_bysiem[element], countalerts, incidents  = generate_eventcount(data, element)
            alldata[element] = incidents
            totalalerts += countalerts
        except:
            print (f"[Error] - No se pudo ejecutar generate_eventcount() en {element}")
            sys.exit()
        all_count_alerts_by_siem[element] = countalerts
        print (f"Total de Alertas en {element}: {countalerts}\n")

        top_patrones_alertas = rule_pattern_analyzer(incidents)
        all_patterns[element] = top_patrones_alertas

        for element2 in all_count_bysiem[element]:
            if element2 not in all_alerts_general_count:
                all_alerts_general_count[element2] = all_count_bysiem[element][element2]
            else:
                all_alerts_general_count[element2] += all_count_bysiem[element][element2]
    
    alerts_by_hour = count_by_hour(alldata)
    
    # FrontEnd Docx
    #______________________________________________________________________________________________________________

    #TOP Grafico General de Alertas
    format = plantilla.add_heading(f"TOP Multiple Fortisiem Count [{totalalerts}]", 1)
    format.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    data_sorted = dict(sorted(all_count_alerts_by_siem.items(), key=lambda x: x[1], reverse=True))
    graficar_gbarras_CLIENTES(data_sorted, "topsiem")
    plantilla.add_picture("topsiem.png", width=docx.shared.Cm(19), height=docx.shared.Cm(10))

    #Tabla General de Alertas
    plantilla.add_heading("Table Multiple Fortisiem Count", 2)
    add_table(plantilla, data_sorted)
    plantilla.add_page_break()


    #TOP Grafico General de alertas
    format = plantilla.add_heading("TOP General Alerts", 1)
    format.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    data_sorted = dict(sorted(all_alerts_general_count.items(), key=lambda x: x[1], reverse=True))

    graficar_gbarras_CLIENTES(data_sorted , "topincidents")
    plantilla.add_picture("topincidents.png", width=docx.shared.Cm(19), height=docx.shared.Cm(10))
    plantilla.add_page_break()

    #TOP Grafico Alertas por Hora
    format = plantilla.add_heading("Alerts by Hour", 1)
    format.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    graficar_gbarras(alerts_by_hour , "alerts_by_hour")
    plantilla.add_picture("alerts_by_hour.png", width=docx.shared.Cm(19), height=docx.shared.Cm(10))
    plantilla.add_page_break()

    plantilla.add_heading("Table General Alerts", 2)
    add_table(plantilla, data_sorted)
    plantilla.add_page_break()

    format = plantilla.add_heading("TOP General Alerts by SIEM", 1)
    format.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for siem_name in all_count_bysiem:
        format = plantilla.add_heading(f"TOP Alerts of {siem_name}", 2)
        format.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        data_sorted = dict(sorted(all_count_bysiem[siem_name].items(), key=lambda x: x[1], reverse=True))
        
        graficar_gbarras_CLIENTES(data_sorted , f"topincidents_{siem_name}")
        plantilla.add_picture(f"topincidents_{siem_name}.png", width=docx.shared.Cm(19), height=docx.shared.Cm(10))
        
        plantilla.add_heading("Top Recurrencia por tipo de alerta [Max=10]", 2)
        add_table(plantilla, all_patterns[siem_name])
        
        plantilla.add_page_break()
   
    #Generar reporte
    fecha = get_time_for_reportname()
    output_docx = (f"output/MFAA_{fecha}.docx")
    
    plantilla.save(output_docx)

    print ("\n\nTotal Alertas:", totalalerts)

main()